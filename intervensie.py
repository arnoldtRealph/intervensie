import streamlit as st
import pandas as pd
import os
from datetime import datetime, timedelta, time
from docx import Document
from docx.shared import Inches
from github import Github
from io import BytesIO
import uuid

# ---------------- Config ---------------- #
st.set_page_config(
    page_title="HOËRSKOOL SAUL DAMON: INTERVENSIE KLASSE",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for modern and smart look with dark mode support
st.markdown(
    """
    <style>
    /* General styling for a modern look */
    body {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        background-color: #f4f7fa;
    }
    .stApp {
        background-color: #ffffff;
        border-radius: 10px;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        padding: 20px;
        max-width: 1200px;
        margin: 0 auto;
    }
    
    /* Dark mode styles */
    @media (prefers-color-scheme: dark) {
        body {
            background-color: #121212;
        }
        .stApp {
            background-color: #1e1e1e;
            color: #e0e0e0;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.3);
        }
        /* Ensure headings are visible in dark mode */
        h1, h2, h3, h4, h5, h6,
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3,
        .stMarkdown h4, .stMarkdown h5, .stMarkdown h6 {
            color: #ffffff !important;
            font-weight: 600 !important;
            text-shadow: 0 1px 2px rgba(0, 0, 0, 0.3) !important;
        }
        .stSelectbox, .stTextInput, .stDateInput, .stTimeInput, .stNumberInput, .stFileUploader {
            background-color: #2d2d2d !important;
            color: #e0e0e0 !important;
            border: 1px solid #404040 !important;
        }
        .stSelectbox > div > div > div {
            color: #e0e0e0 !important;
        }
        .stButton > button {
            background-color: #007bff !important;
            color: white !important;
            border-radius: 8px !important;
            padding: 10px 20px !important;
            font-weight: 500 !important;
            transition: background-color 0.3s ease !important;
        }
        .stButton > button:hover {
            background-color: #0056b3 !important;
        }
        div.stButton > button[kind="formSubmit"] {
            background-color: #28a745 !important;
            color: white !important;
            border: none !important;
            padding: 0.5rem 1rem !important;
            border-radius: 8px !important;
            width: 100% !important;
            font-weight: 500 !important;
        }
        div.stButton > button[kind="formSubmit"]:hover {
            background-color: #218838 !important;
        }
        .stSidebar {
            background-color: #2a2a2a !important;
            color: #e0e0e0 !important;
            border-radius: 10px !important;
            padding: 15px !important;
        }
        .stSidebar .stSelectbox, .stSidebar .stTextInput {
            background-color: #3a3a3a !important;
            color: #e0e0e0 !important;
            border: 1px solid #505050 !important;
        }
        .stDataFrame {
            background-color: #2d2d2d !important;
            color: #e0e0e0 !important;
            border-radius: 8px !important;
            overflow: hidden !important;
        }
        .stDataFrame thead tr th {
            background-color: #404040 !important;
            color: #ffffff !important;
        }
        .stDataFrame tbody tr td {
            background-color: #2d2d2d !important;
            color: #e0e0e0 !important;
            border-bottom: 1px solid #404040 !important;
        }
        .stDataFrame tbody tr:nth-child(even) td {
            background-color: #333333 !important;
        }
        .stCaption {
            color: #b0b0b0 !important;
            font-style: italic !important;
        }
        .stInfo, .stWarning, .stError, .stSuccess {
            border-radius: 8px !important;
        }
        /* Mobile specific adjustments */
        @media (max-width: 768px) {
            .stApp {
                padding: 10px !important;
                margin: 5px !important;
                border-radius: 5px !important;
            }
            .stSidebar {
                padding: 10px !important;
                border-radius: 5px !important;
            }
            /* Ensure headings remain visible on mobile in dark mode */
            h1, h2, h3, h4, h5, h6,
            .stMarkdown h1, .stMarkdown h2, .stMarkdown h3,
            .stMarkdown h4, .stMarkdown h5, .stMarkdown h6 {
                color: #ffffff !important;
            }
        }
    }
    
    /* Light mode styles (default) */
    @media (prefers-color-scheme: light) {
        h1, h2, h3, h4, h5, h6,
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3,
        .stMarkdown h4, .stMarkdown h5, .stMarkdown h6 {
            color: #343a40 !important;
            font-weight: 600 !important;
            text-shadow: 0 1px 2px rgba(0, 0, 0, 0.1) !important;
        }
        .stSelectbox, .stTextInput, .stDateInput, .stTimeInput, .stNumberInput, .stFileUploader {
            background-color: #f8f9fa !important;
            border-radius: 8px !important;
            padding: 10px !important;
            border: 1px solid #ced4da !important;
            color: #343a40 !important;
        }
        .stButton > button {
            background-color: #007bff !important;
            color: white !important;
            border-radius: 8px !important;
            padding: 10px 20px !important;
            font-weight: 500 !important;
            transition: background-color 0.3s ease !important;
        }
        .stButton > button:hover {
            background-color: #0056b3 !important;
        }
        .stSidebar {
            background-color: #e9ecef !important;
            border-radius: 10px !important;
            padding: 15px !important;
        }
        .stDataFrame {
            border-radius: 8px !important;
            overflow: hidden !important;
        }
        .stCaption {
            color: #6c757d !important;
            font-style: italic !important;
        }
        /* Mobile specific adjustments */
        @media (max-width: 768px) {
            .stApp {
                padding: 10px !important;
                margin: 5px !important;
                border-radius: 5px !important;
            }
            .stSidebar {
                padding: 10px !important;
                border-radius: 5px !important;
            }
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Constants
CSV_FILE = "intervensie_database.csv"
LOG_FILE = "app_log.csv"
ERROR_LOG_FILE = "error_log.txt"
FOTO_DIR = "fotos"
PRES_DIR = "presensies"
GRADE_OPTIONS = ["8", "9", "10", "11", "12"]

# Initialize directories and CSV
for directory in [FOTO_DIR, PRES_DIR]:
    os.makedirs(directory, exist_ok=True)

if not os.path.exists(CSV_FILE):
    pd.DataFrame(columns=[
        "Datum", "Graad", "Vak", "Tema", "Begintyd", "Eindtyd", 
        "Totaal Genooi", "Totaal Opgedaag", "Opvoeder", "Foto", 
        "Presensielys_Foto", "Presensielys_Dokument"
    ]).to_csv(CSV_FILE, index=False)

if not os.path.exists(LOG_FILE):
    pd.DataFrame(columns=["Timestamp", "Action", "Details", "Status"]).to_csv(LOG_FILE, index=False)

# ---------------- Log Functions ---------------- #
def log_action(action, details="", status="INFO"):
    """Log actions to CSV file."""
    log_entry = {
        "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Action": action,
        "Details": details,
        "Status": status
    }
    try:
        df_log = pd.read_csv(LOG_FILE)
        df_log = pd.concat([df_log, pd.DataFrame([log_entry])], ignore_index=True)
        df_log.to_csv(LOG_FILE, index=False)
        return True
    except Exception as e:
        st.error(f"⚠️ Fout met log stoor: {str(e)}")
        with open(ERROR_LOG_FILE, "a") as f:
            f.write(f"Log save failed: {str(e)} at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        return False

# ---------------- GitHub Upload Function ---------------- #
def upload_file_to_github(file_path, repo_name, path_in_repo, token):
    """Upload or update a file in a GitHub repository using PyGithub."""
    try:
        log_action("GitHub Upload Attempt", f"File: {path_in_repo}, Repo: {repo_name}", "INFO")
        if not token or token.strip() == "":
            log_action("GitHub Upload Failed", "Empty or missing token", "ERROR")
            st.error("⚠️ GitHub token is leeg of ontbreek!")
            return False

        g = Github(token)
        repo = g.get_repo(repo_name)

        if not os.path.exists(file_path):
            log_action("GitHub Upload Failed", f"Local file not found: {file_path}", "ERROR")
            st.error(f"⚠️ Lokale lêer nie gevind nie: {file_path}")
            return False

        with open(file_path, "rb") as file:
            content = file.read()

        repo_path = path_in_repo
        try:
            contents = repo.get_contents(repo_path, ref="master")
            repo.update_file(
                path=repo_path,
                message=f"Updated {repo_path} - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                content=content,
                sha=contents.sha,
                branch="master"
            )
            log_action("GitHub Upload Success", f"Updated existing file: {repo_path}", "SUCCESS")
        except Exception:
            repo.create_file(
                path=repo_path,
                message=f"Created {repo_path} - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                content=content,
                branch="master"
            )
            log_action("GitHub Upload Success", f"Created new file: {repo_path}", "SUCCESS")
        return True
    except Exception as e:
        error_msg = str(e)
        log_action("GitHub Upload Failed", f"Error: {error_msg}", "ERROR")
        with open(ERROR_LOG_FILE, "a") as f:
            f.write(f"GitHub push failed: {error_msg} at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        st.error(f"⚠️ GitHub upload misluk: {error_msg}")
        return False

# ---------------- Helper: safe read attendance file ---------------- #
def read_presensie_to_table(path, max_rows=50):
    """Try to convert a CSV/XLSX presensielys into a pandas DataFrame for insertion into Word."""
    try:
        ext = path.split('.')[-1].lower()
        if ext == 'csv':
            df_p = pd.read_csv(path)
        elif ext in ['xls', 'xlsx']:
            df_p = pd.read_excel(path)
        else:
            return None
        if df_p.shape[0] > max_rows:
            return df_p.iloc[:max_rows]
        return df_p
    except Exception as e:
        log_action("Presensie Read Failed", f"{path} - {str(e)}", "WARNING")
        return None

# ---------------- Load Intervention Data ---------------- #
@st.cache_data(ttl=600)
def load_intervention_data():
    if not os.path.exists(CSV_FILE):
        return pd.DataFrame()
    df = pd.read_csv(CSV_FILE)
    if df.empty:
        return df
    df["Datum"] = pd.to_datetime(df["Datum"], errors="coerce")
    df["Aanwesigheid %"] = (df["Totaal Opgedaag"] / df["Totaal Genooi"] * 100).round(2)
    return df.sort_values("Datum", ascending=False)

# ---------------- Load Raw Data for Filters and Deletion ---------------- #
@st.cache_data(ttl=300)
def load_raw():
    if not os.path.exists(CSV_FILE):
        return pd.DataFrame()
    df = pd.read_csv(CSV_FILE)
    if df.empty:
        return df
    df['Datum'] = pd.to_datetime(df['Datum'], errors='coerce')
    return df.sort_values("Datum", ascending=False)

# ---------------- UI ---------------- #
st.title("HOËRSKOOL SAUL DAMON")
st.subheader("📘 Intervensie Klasse")

# Sidebar filters for Word report
st.sidebar.header("Filters vir Verslag")
filter_type = st.sidebar.selectbox("🔎 Kies tydsfilter", ["Alles", "Weekliks", "Maandeliks", "Kwartaalliks", "Jaarliks"]) 

raw_df = load_raw()

# Options for filter selectors
opvoeder_options = ['Alles'] + sorted(raw_df['Opvoeder'].dropna().unique().tolist()) if not raw_df.empty else ['Alles']
vak_options = ['Alles'] + sorted(raw_df['Vak'].dropna().unique().tolist()) if not raw_df.empty else ['Alles']
graad_options = ['Alles'] + GRADE_OPTIONS

selected_opvoeder = st.sidebar.selectbox("Opvoeder", opvoeder_options)
selected_vak = st.sidebar.selectbox("Vak", vak_options)
selected_graad = st.sidebar.selectbox("Graad", graad_options)

# Form for new entries
with st.form("data_form", clear_on_submit=True):
    col1, col2 = st.columns(2)
    with col1:
        datum = st.date_input("📅 Datum", value=datetime.today(), format="YYYY/MM/DD")
        graad = st.selectbox("🎓 Graad", GRADE_OPTIONS, key='form_graad')
        vak = st.text_input("📚 Vak", key='form_vak')
        tema = st.text_input("🎯 Tema", key='form_tema')
        begintyd = st.time_input("🕒 Begintyd", value=time(8, 0), step=900)  # 15-minute intervals
        eindtyd = st.time_input("🕔 Eindtyd", value=time(9, 0), step=900)    # 15-minute intervals
    with col2:
        totaal_genooi = st.number_input("👥 Totaal Genooi", min_value=1, step=1, format="%d", key='form_totaal_genooi')
        totaal_opgedaag = st.number_input("✅ Totaal Opgedaag", min_value=0, step=1, format="%d", key='form_totaal_opgedaag')
        opvoeder = st.text_input("👨‍🏫 Opvoeder", key='form_opvoeder')
    
    st.subheader("📸 Presensielys Upload")
    presensie_foto = st.file_uploader(
        "📷 Laai Presensielys Foto op (opsioneel)", 
        type=["jpg", "jpeg", "png"],
        key='form_presensie_foto'
    )
    presensie_dokument = st.file_uploader(
        "📑 Laai Presensielys Dokument op (opsioneel)", 
        type=["csv", "xlsx", "pdf"],
        key='form_presensie_dokument'
    )
    foto = st.file_uploader("📸 Laai Foto op", type=["jpg", "jpeg", "png"], key='form_foto')

    # Submit button
    submitted = st.form_submit_button(
        "➕ Stoor Intervensie",
        help="Stoor die intervensie data",
        use_container_width=True
    )

    if submitted:
        log_action("Form Submission", f"Submitted by: {opvoeder}", "INFO")
        if not all([datum, graad, vak, tema, begintyd, eindtyd, opvoeder, totaal_genooi]):
            log_action("Form Validation Failed", "Missing required fields", "WARNING")
            st.error("⚠️ Alle verpligte velde (behalwe presensielys) moet ingevul word!")
        elif totaal_opgedaag > totaal_genooi:
            log_action("Form Validation Failed", f"Attendance ({totaal_opgedaag}) > Total ({totaal_genooi})", "WARNING")
            st.error("⚠️ Totaal Opgedaag kan nie meer as Totaal Genooi wees nie!")
        elif begintyd >= eindtyd:
            log_action("Form Validation Failed", f"Start time ({begintyd}) >= End time ({eindtyd})", "WARNING")
            st.error("⚠️ Eindtyd moet later as Begintyd wees!")
        elif not (presensie_foto or presensie_dokument):
            log_action("Form Validation Failed", "No presensielys uploaded", "WARNING")
            st.error("⚠️ Ten minste een presensielys (foto of dokument) moet opgelaai word!")
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            foto_path = ""
            pres_foto_path = ""
            pres_dokument_path = ""

            if foto:
                foto_ext = os.path.splitext(foto.name)[1]
                foto_path = os.path.join(FOTO_DIR, f"foto_{timestamp}{foto_ext}")
                try:
                    with open(foto_path, "wb") as f:
                        f.write(foto.getbuffer())
                    log_action("File Save Success", f"Photo saved: {foto_path}", "SUCCESS")
                except Exception as e:
                    log_action("File Save Failed", f"Photo save error: {str(e)}", "ERROR")
                    st.error(f"⚠️ Fout met foto stoor: {str(e)}")
                    st.stop()

            if presensie_foto:
                pres_foto_ext = os.path.splitext(presensie_foto.name)[1]
                pres_foto_path = os.path.join(PRES_DIR, f"presensie_foto_{timestamp}{pres_foto_ext}")
                try:
                    with open(pres_foto_path, "wb") as f:
                        f.write(presensie_foto.getbuffer())
                    log_action("File Save Success", f"Presensie foto saved: {pres_foto_path}", "SUCCESS")
                except Exception as e:
                    log_action("File Save Failed", f"Presensie foto save error: {str(e)}", "ERROR")
                    st.error(f"⚠️ Fout met presensielys foto stoor: {str(e)}")
                    st.stop()

            if presensie_dokument:
                pres_dokument_ext = os.path.splitext(presensie_dokument.name)[1]
                pres_dokument_path = os.path.join(PRES_DIR, f"presensie_dokument_{timestamp}{pres_dokument_ext}")
                try:
                    with open(pres_dokument_path, "wb") as f:
                        f.write(presensie_dokument.getbuffer())
                    log_action("File Save Success", f"Presensie dokument saved: {pres_dokument_path}", "SUCCESS")
                except Exception as e:
                    log_action("File Save Failed", f"Presensie dokument save error: {str(e)}", "ERROR")
                    st.error(f"⚠️ Fout met presensielys dokument stoor: {str(e)}")
                    st.stop()

            try:
                new_entry = {
                    "Datum": datum.strftime("%Y-%m-%d"),
                    "Graad": graad,
                    "Vak": vak,
                    "Tema": tema,
                    "Begintyd": begintyd.strftime("%H:%M"),
                    "Eindtyd": eindtyd.strftime("%H:%M"),
                    "Totaal Genooi": int(totaal_genooi),
                    "Totaal Opgedaag": int(totaal_opgedaag),
                    "Opvoeder": opvoeder,
                    "Foto": foto_path,
                    "Presensielys_Foto": pres_foto_path,
                    "Presensielys_Dokument": pres_dokument_path
                }
                df = pd.read_csv(CSV_FILE)
                df = pd.concat([df, pd.DataFrame([new_entry])], ignore_index=True)
                df.to_csv(CSV_FILE, index=False)
                log_action("Database Update Success", f"Added entry for {datum.strftime('%Y-%m-%d')} - {vak}", "SUCCESS")
            except Exception as e:
                log_action("Database Update Failed", f"CSV error: {str(e)}", "ERROR")
                st.error(f"⚠️ Fout met databasis stoor: {str(e)}")
                st.stop()

            try:
                token = st.secrets.get("GITHUB_TOKEN")
                repo = st.secrets.get("GITHUB_REPO")
                if not token or not repo:
                    log_action("GitHub Config Missing", f"Token: {bool(token)}, Repo: {bool(repo)}", "WARNING")
                    st.warning("⚠️ GitHub konfigurasie ontbreek in secrets.")
                elif upload_file_to_github(CSV_FILE, repo, "intervensie_database.csv", token):
                    log_action("Sync Complete", "All operations successful", "SUCCESS")
                    st.success("✅ Data gestoor en gesinkroniseer met GitHub!")
                else:
                    log_action("Sync Incomplete", "GitHub sync failed but data saved locally", "WARNING")
                    st.warning("⚠️ Data lokaal gestoor, maar GitHub sinkronisasie misluk.")
            except KeyError as e:
                log_action("GitHub Secrets Error", f"Missing secret: {str(e)}", "ERROR")
                st.error("⚠️ GitHub konfigurasie ontbreek in secrets!")
            except Exception as e:
                log_action("GitHub Unexpected Error", f"Sync error: {str(e)}", "ERROR")
                st.error(f"⚠️ Onverwagte GitHub fout: {str(e)}")

            # Clear cache and rerun to update log display immediately
            load_intervention_data.clear()
            load_raw.clear()
            st.rerun()

# ---------------- Log Display (Intervention Data) ---------------- #
st.subheader("📊 Intervensie Log Inskrywings")

intervention_df = load_intervention_data()

if 'intervention_page' not in st.session_state:
    st.session_state.intervention_page = 0

ENTRIES_PER_PAGE = 10
total_entries = len(intervention_df)
total_pages = (total_entries + ENTRIES_PER_PAGE - 1) // ENTRIES_PER_PAGE

# Calculate start and end indices for current page
start_idx = st.session_state.intervention_page * ENTRIES_PER_PAGE
end_idx = min(start_idx + ENTRIES_PER_PAGE, total_entries)

# Display intervention data on homepage
if intervention_df.empty:
    st.info("ℹ️ Geen intervensie inskrywings nie.")
else:
    log_action("Intervention Log Report Generated", f"Records: {len(intervention_df)}", "INFO")
    st.dataframe(
        intervention_df.iloc[start_idx:end_idx][["Datum", "Graad", "Vak", "Tema", "Begintyd", "Eindtyd", "Totaal Genooi", "Totaal Opgedaag", "Opvoeder", "Aanwesigheid %"]].reset_index(drop=True),
        column_config={
            "Datum": st.column_config.DateColumn(format="YYYY-MM-DD"),
            "Aanwesigheid %": st.column_config.NumberColumn(format="%.2f%%"),
            "Graad": st.column_config.SelectboxColumn(options=GRADE_OPTIONS),
            "Begintyd": st.column_config.TextColumn(),
            "Eindtyd": st.column_config.TextColumn()
        },
        use_container_width=True
    )

    # Pagination controls
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.session_state.intervention_page > 0:
            if st.button("Vorige"):
                st.session_state.intervention_page -= 1
                st.rerun()  # Rerun to update pagination
    with col3:
        if st.session_state.intervention_page < total_pages - 1:
            if st.button("Volgende"):
                st.session_state.intervention_page += 1
                st.rerun()  # Rerun to update pagination
    with col2:
        st.write(f"Bladsy {st.session_state.intervention_page + 1} van {max(total_pages,1)}")

# ---------------- Load and Filter Intervention Data for Report and Deletion ---------------- #
@st.cache_data(ttl=600)
def load_and_filter_data(filter_type, opvoeder=None, vak=None, graad=None):
    if not os.path.exists(CSV_FILE):
        return pd.DataFrame()
    df = pd.read_csv(CSV_FILE)
    if df.empty:
        return df
    df["Datum"] = pd.to_datetime(df["Datum"], errors="coerce")
    df["Aanwesigheid %"] = (df["Totaal Opgedaag"] / df["Totaal Genooi"] * 100).round(2)

    today = datetime.today()
    if filter_type == "Weekliks":
        start = today - timedelta(days=7)
        df = df[df["Datum"] >= start]
    elif filter_type == "Maandeliks":
        start = today - timedelta(days=30)
        df = df[df["Datum"] >= start]
    elif filter_type == "Kwartaalliks":
        start = today - timedelta(days=90)
        df = df[df["Datum"] >= start]
    elif filter_type == "Jaarliks":
        start = today - timedelta(days=365)
        df = df[df["Datum"] >= start]

    # Apply additional filters
    if opvoeder and opvoeder != 'Alles':
        df = df[df['Opvoeder'] == opvoeder]
    if vak and vak != 'Alles':
        df = df[df['Vak'] == vak]
    if graad and graad != 'Alles':
        df = df[df['Graad'] == graad]

    return df.sort_values("Datum", ascending=False)

# Load filtered data for Word report
df = load_and_filter_data(filter_type, selected_opvoeder, selected_vak, selected_graad)

# ---------------- Deletion ---------------- #
st.subheader("🗑️ Verwyder Intervensie Inskrywing")
if not raw_df.empty:
    entries = [f"ID {idx}: {row['Datum'].strftime('%Y-%m-%d')} - {row['Vak']} - {row['Opvoeder']}" for idx, row in raw_df.iterrows()]
    selected_entry = st.selectbox("Kies inskrywing om te verwyder", ["Geen"] + entries)
    if st.button("Bevestig Verwydering"):
        if selected_entry != "Geen":
            try:
                idx = int(selected_entry.split(":")[0].split(" ")[1])
                full_df = pd.read_csv(CSV_FILE)
                row_to_delete = full_df.loc[idx]
                full_df = full_df.drop(idx).reset_index(drop=True)
                full_df.to_csv(CSV_FILE, index=False)

                # Delete associated files
                if pd.notna(row_to_delete['Foto']) and os.path.exists(row_to_delete['Foto']):
                    os.remove(row_to_delete['Foto'])
                    log_action("File Delete Success", f"Photo deleted: {row_to_delete['Foto']}", "SUCCESS")
                if pd.notna(row_to_delete['Presensielys_Foto']) and os.path.exists(row_to_delete['Presensielys_Foto']):
                    os.remove(row_to_delete['Presensielys_Foto'])
                    log_action("File Delete Success", f"Presensielys foto deleted: {row_to_delete['Presensielys_Foto']}", "SUCCESS")
                if pd.notna(row_to_delete['Presensielys_Dokument']) and os.path.exists(row_to_delete['Presensielys_Dokument']):
                    os.remove(row_to_delete['Presensielys_Dokument'])
                    log_action("File Delete Success", f"Presensielys dokument deleted: {row_to_delete['Presensielys_Dokument']}", "SUCCESS")

                # Sync to GitHub
                token = st.secrets.get("GITHUB_TOKEN")
                repo = st.secrets.get("GITHUB_REPO")
                if token and repo:
                    upload_file_to_github(CSV_FILE, repo, "intervensie_database.csv", token)

                st.success("✅ Inskrywing suksesvol verwyder!")
                log_action("Deletion Success", f"Deleted ID {idx}", "SUCCESS")
                load_and_filter_data.clear()
                load_raw.clear()
                load_intervention_data.clear()
                st.rerun()  # Rerun to update log display after deletion
            except Exception as e:
                st.error(f"⚠️ Fout met verwydering: {str(e)}")
                log_action("Deletion Failed", f"Error: {str(e)}", "ERROR")
else:
    st.info("ℹ️ Geen inskrywings beskikbaar om te verwyder nie.")

# ---------------- Word Report Download ---------------- #
st.subheader("📑 Intervensie Verslag Aflaai")

def generate_word_report(df_to_export):
    doc = Document()
    doc.add_heading("Saul Damon High School - Intervensie Verslag", level=1)
    doc.add_paragraph(f"Filter: {filter_type} | Opvoeder: {selected_opvoeder} | Vak: {selected_vak} | Graad: {selected_graad}")
    doc.add_paragraph(f"Gegenereer op: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    if not df_to_export.empty:
        # Summary table
        columns = ["Datum", "Graad", "Vak", "Tema", "Begintyd", "Eindtyd", "Totaal Genooi", "Totaal Opgedaag", "Opvoeder", "Aanwesigheid %"]
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = col

        for _, row in df_to_export.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row['Datum'].strftime('%Y-%m-%d')
            row_cells[1].text = str(row['Graad'])
            row_cells[2].text = row['Vak']
            row_cells[3].text = row['Tema']
            row_cells[4].text = str(row.get('Begintyd', 'NVT'))
            row_cells[5].text = str(row.get('Eindtyd', 'NVT'))
            row_cells[6].text = str(row['Totaal Genooi'])
            row_cells[7].text = str(row['Totaal Opgedaag'])
            row_cells[8].text = row['Opvoeder']
            row_cells[9].text = f"{row['Aanwesigheid %']:.2f}%"

        doc.add_paragraph("")
        doc.add_heading("Details met Fotos en Presensielyste", level=2)

        for _, row in df_to_export.iterrows():
            doc.add_heading(f"Inskrywing: {row['Datum'].strftime('%Y-%m-%d')} - {row['Vak']} - {row.get('Begintyd', 'NVT')} tot {row.get('Eindtyd', 'NVT')}", level=3)

            # Foto insertion
            if pd.notna(row.get('Foto')) and os.path.exists(row['Foto']):
                try:
                    doc.add_paragraph('Foto:')
                    doc.add_picture(row['Foto'], width=Inches(2))
                except Exception as e:
                    doc.add_paragraph(f"⚠️ Kon nie foto laai nie: {str(e)}")
            else:
                doc.add_paragraph("Geen geldige foto gevind nie.")

            # Presensielys Foto insertion
            doc.add_paragraph('Presensielys Foto:')
            if pd.notna(row.get('Presensielys_Foto')) and os.path.exists(row['Presensielys_Foto']):
                try:
                    doc.add_picture(row['Presensielys_Foto'], width=Inches(2))
                except Exception as e:
                    doc.add_paragraph(f"⚠️ Kon nie presensielys foto laai nie: {str(e)}")
            else:
                doc.add_paragraph("Geen presensielys foto opgelaai nie.")

            # Presensielys Dokument handling
            doc.add_paragraph('Presensielys Dokument:')
            if pd.notna(row.get('Presensielys_Dokument')) and os.path.exists(row['Presensielys_Dokument']):
                pres_path = row['Presensielys_Dokument']
                ext = pres_path.split('.')[-1].lower()
                if ext in ['csv', 'xls', 'xlsx']:
                    df_p = read_presensie_to_table(pres_path)
                    if df_p is not None and not df_p.empty:
                        sub_table = doc.add_table(rows=1, cols=min(len(df_p.columns), 10))
                        sub_hdr_cells = sub_table.rows[0].cells
                        for i, col_name in enumerate(df_p.columns[:10]):
                            sub_hdr_cells[i].text = str(col_name)
                        for _, prow in df_p.iterrows():
                            sub_row_cells = sub_table.add_row().cells
                            for i, val in enumerate(prow[:10]):
                                sub_row_cells[i].text = str(val)
                        if len(df_p) >= 50:
                            doc.add_paragraph('... (tabel afgekort — slegs die eerste rye getoon)')
                    else:
                        doc.add_paragraph(f"Kon nie presensielys lees nie: {os.path.basename(pres_path)}")
                else:
                    doc.add_paragraph(f"Lêer: {os.path.basename(pres_path)} (PDF of onbekende tipe - word in die map gehou)")
            else:
                doc.add_paragraph("Geen presensielys dokument opgelaai nie.")

            doc.add_paragraph("-" * 30)
    else:
        doc.add_paragraph("Geen data vir die gekose filters nie.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Download button for Word report
try:
    doc_bytes = generate_word_report(df)
    st.download_button(
        label="⬇️ Laai Intervensie Verslag af (Word)",
        data=doc_bytes,
        file_name=f"intervensie_report_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_word_report"
    )
except Exception as e:
    log_action("Word Report Download Failed", f"Error: {str(e)}", "ERROR")
    st.error(f"⚠️ Fout met verslag aflaai: {str(e)}")

# Small note for users about large presensielyste
st.caption("Let asseblief: Groter presensielyste (baie rye) word afgekort in die Word verslag om dokumentgrootte te beperk. Indien nodig, laai die oorspronklike lêer af vanaf die server se 'presensies' gids.")
